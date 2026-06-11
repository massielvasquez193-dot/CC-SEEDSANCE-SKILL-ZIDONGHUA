"""Generate quotation.docx for TikTok AI Factory Pro"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('TikTok AI Factory Pro', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI Video Production Factory - Quotation')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Date: {datetime.now().strftime("%Y-%m-%d")}  |  Version: v3.0.0')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# Overview
doc.add_heading('Product Overview', level=1)
overview = doc.add_paragraph()
overview.add_run('TikTok AI Factory Pro ').bold = True
overview.add_run(
    'is a fully automated TikTok video production system. '
    'Simply place product images, reference videos, and character images into the input folder. '
    'The system automatically handles: GPT-5.5 script generation, GPT Image keyframe creation, '
    'Seedance 2.0 video generation, ElevenLabs voiceover, FFmpeg compositing, and final video output.'
)

doc.add_paragraph()

# Version Comparison Table
doc.add_heading('Version Comparison', level=1)

table = doc.add_table(rows=15, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Feature', 'Basic', 'Professional', 'Enterprise']
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = cell._element.get_or_add_tcPr()
    shading.append(shading.makeelement(qn('w:shd'), {qn('w:fill'): '1A1A2E', qn('w:val'): 'clear'}))

# Price row
prices = ['', 'CNY 999/mo', 'CNY 1,999/mo', 'CNY 4,999/mo']
for i, price in enumerate(prices):
    if i == 0:
        table.rows[1].cells[i].text = 'Price'
        table.rows[1].cells[i].paragraphs[0].runs[0].bold = True
    else:
        cell = table.rows[1].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(price)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

# Feature rows
features = [
    ('Installation & Deployment', ['YES', 'YES', 'YES']),
    ('Environment Configuration', ['YES', 'YES', 'YES']),
    ('Basic Script Generation', ['YES', 'YES', 'YES']),
    ('Basic Storyboard', ['YES', 'YES', 'YES']),
    ('GPT-5.5 AI Script', ['--', 'YES', 'YES']),
    ('GPT Image Keyframes', ['--', 'YES', 'YES']),
    ('Seedance 2.0 Video Gen', ['--', 'YES', 'YES']),
    ('ElevenLabs Voiceover', ['--', 'YES', 'YES']),
    ('FFmpeg Compositing', ['--', 'YES', 'YES']),
    ('Multi-Account Support', ['--', '--', 'YES']),
    ('Batch Production Mode', ['--', '--', 'YES']),
    ('Watch Mode (Auto-detect)', ['--', '--', 'YES']),
    ('Commercial License System', ['--', '--', 'YES']),
]

for row_idx, (feature, values) in enumerate(features):
    row = table.rows[row_idx + 2]
    row.cells[0].text = feature
    for col_idx, val in enumerate(values):
        cell = row.cells[col_idx + 1]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        if val == 'YES':
            run.font.color.rgb = RGBColor(0x23, 0x86, 0x36)
        else:
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

doc.add_paragraph()

# Version Details
doc.add_heading('Version Details', level=1)

# Basic
doc.add_heading('Basic - CNY 999/month', level=2)
p = doc.add_paragraph()
p.add_run('For: ').bold = True
p.add_run('Individual creators exploring AI video production')
for item in ['Python environment setup & deployment', 'Project structure configuration',
             '.env API key setup guide', 'Basic script generation (template mode)',
             'Basic storyboard generation', 'Basic subtitle generation',
             'Summary reports', '7-day technical support']:
    doc.add_paragraph(item, style='List Bullet')

# Professional
doc.add_heading('Professional - CNY 1,999/month', level=2)
p = doc.add_paragraph()
p.add_run('For: ').bold = True
p.add_run('Professional content creators, MCN agencies')
for item in ['Everything in Basic', 'GPT-5.5 AI script generation (UGC style)',
             'GPT Image real-person keyframes (no placeholders)',
             'Seedance 2.0 ARK API video generation',
             'ElevenLabs real-person TTS voiceover',
             'FFmpeg video + voice + subtitle compositing',
             'Character consistency engine', 'Product consistency engine',
             'UGC quality scoring system', 'Performance dashboard',
             '30-day technical support']:
    doc.add_paragraph(item, style='List Bullet')

# Enterprise
doc.add_heading('Enterprise - CNY 4,999/month', level=2)
p = doc.add_paragraph()
p.add_run('For: ').bold = True
p.add_run('Batch operators, cross-border e-commerce enterprises')
for item in ['Everything in Professional', 'Multi-account management',
             'Batch production (unlimited tasks)', 'Watch Mode auto-processing',
             'Commercial license system (hardware-locked)',
             'Multi-character library (US/UK/AU)', 'Multi-country/language (13 countries)',
             'UGC Director performance engine', 'Batch task dashboard',
             'Priority technical support', 'Custom feature development (on demand)',
             'API quota monitoring & alerts']:
    doc.add_paragraph(item, style='List Bullet')

# Tech Stack
doc.add_heading('Technology Stack', level=1)
doc.add_paragraph(
    'GPT-5.5 (Script) | GPT Image/DALL-E 3 (Keyframes) | Seedance 2.0 (Video) | '
    'ElevenLabs (TTS) | FFmpeg (Compositing) | '
    'OpenAI API | Anthropic Claude | Google Gemini | DeepSeek | '
    'Volcengine ARK | Runway | Kling'
)

# Delivery & Payment
doc.add_heading('Delivery & Payment', level=1)
p = doc.add_paragraph()
p.add_run('Delivery: ').bold = True
p.add_run('GitHub private repository + deployment documentation + video tutorial')
p = doc.add_paragraph()
p.add_run('Payment: ').bold = True
p.add_run('Bank transfer / Alipay / WeChat Pay / USDT')
p = doc.add_paragraph()
p.add_run('Invoice: ').bold = True
p.add_run('VAT invoice available')
p = doc.add_paragraph()
p.add_run('Trial: ').bold = True
p.add_run('3-day free trial (Basic features)')

# Contact
doc.add_heading('Contact', level=1)
doc.add_paragraph('GitHub: https://github.com/massielvasquez193-dot/CC-SEEDSANCE-SKILL-ZIDONGHUA')

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('TikTok AI Factory Pro - AI Video Production Factory')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer2.add_run(f'Quote valid until {datetime.now().year}-12-31 | Prices exclude tax | Terms apply')
run2.font.size = Pt(8)
run2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

doc.save('报价单.docx')
print('Done: 报价单.docx')
