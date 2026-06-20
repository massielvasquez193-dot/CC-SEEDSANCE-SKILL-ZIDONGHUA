"""Generate quotation CN version"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.size = Pt(11)

# === TITLE ===
title = doc.add_heading('TikTok AI Factory Pro', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('全自动AI视频生产工厂 · 产品报价单')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = d.add_run(f'报价日期: {datetime.now().strftime("%Y年%m月%d日")}  |  版本: v3.0.0  |  币种: 人民币 (CNY)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# === OVERVIEW ===
doc.add_heading('一、产品概述', level=1)
p = doc.add_paragraph()
p.add_run('TikTok AI Factory Pro ').bold = True
p.add_run(
    '是一套全自动TikTok视频生产系统。无需编程知识，只需将产品图片、参考视频和人物图片放入指定文件夹，'
    '系统即可自动完成以下全流程：GPT-5.5智能脚本撰写 → GPT Image真人关键帧生成 → Seedance 2.0视频生成 → '
    'ElevenLabs真人配音 → FFmpeg视频合成 → 最终视频输出。'
)
p = doc.add_paragraph()
p.add_run('核心技术栈: ').bold = True
p.add_run('GPT-5.5 脚本引擎 | GPT Image/DALL-E 3 关键帧 | Seedance 2.0 视频生成 | ElevenLabs TTS | FFmpeg 合成')

doc.add_paragraph()

# === VERSION TABLE ===
doc.add_heading('二、版本对比', level=1)

table = doc.add_table(rows=15, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['功能模块', '基础版', '专业版', '企业版']
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = cell._element.get_or_add_tcPr()
    shading.append(shading.makeelement(qn('w:shd'), {qn('w:fill'): '1A1A2E', qn('w:val'): 'clear'}))

prices = ['', '¥999/月', '¥1,999/月', '¥4,999/月']
for i, price in enumerate(prices):
    if i == 0:
        table.rows[1].cells[i].text = '价格'
        table.rows[1].cells[i].paragraphs[0].runs[0].bold = True
    else:
        cell = table.rows[1].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(price)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

features = [
    ('安装与部署', ['✓', '✓', '✓']),
    ('环境配置指导', ['✓', '✓', '✓']),
    ('基础脚本生成', ['✓', '✓', '✓']),
    ('基础分镜表', ['✓', '✓', '✓']),
    ('GPT-5.5 AI智能脚本', ['—', '✓', '✓']),
    ('GPT Image 真人关键帧', ['—', '✓', '✓']),
    ('Seedance 2.0 视频生成', ['—', '✓', '✓']),
    ('ElevenLabs 真人配音', ['—', '✓', '✓']),
    ('FFmpeg 视频合成', ['—', '✓', '✓']),
    ('多账号管理', ['—', '—', '✓']),
    ('批量生产模式', ['—', '—', '✓']),
    ('Watch Mode 自动监控', ['—', '—', '✓']),
    ('商业授权系统', ['—', '—', '✓']),
]

for row_idx, (feature, values) in enumerate(features):
    row = table.rows[row_idx + 2]
    row.cells[0].text = feature
    for col_idx, val in enumerate(values):
        cell = row.cells[col_idx + 1]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(12)
        if val == '✓':
            run.font.color.rgb = RGBColor(0x23, 0x86, 0x36)
        else:
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

doc.add_paragraph()

# === VERSION DETAILS ===
doc.add_heading('三、版本详情', level=1)

# Basic
doc.add_heading('基础版 — ¥999/月', level=2)
p = doc.add_paragraph()
p.add_run('适合: ').bold = True
p.add_run('个人创作者、初步尝试AI视频生产的用户')
items = [
    'Python运行环境安装与部署',
    '项目目录结构搭建与配置',
    '.env API密钥配置指导',
    '基础视频脚本生成 (模板模式)',
    '基础分镜表生成 (Storyboard)',
    '基础字幕文件生成 (SRT格式)',
    '基础任务总结报告',
    '7天微信/邮件技术支持',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Pro
doc.add_heading('专业版 — ¥1,999/月', level=2)
p = doc.add_paragraph()
p.add_run('适合: ').bold = True
p.add_run('专业内容创作者、MCN机构、电商运营团队')
items = [
    '基础版全部功能',
    'GPT-5.5 智能脚本生成 (真人UGC口播风格, 6段式结构)',
    'GPT Image (DALL-E 3) 真人关键帧生成 (禁止占位图, 真实美国家庭场景)',
    'Seedance 2.0 ARK API 视频生成 (火山引擎, 支持图生视频)',
    'ElevenLabs 真人TTS配音 (8种情绪映射, 29种语言)',
    'FFmpeg 视频+配音+字幕合成 (UGC字幕样式)',
    '人物一致性引擎 (全视频同人同发型同服装同妆容)',
    '产品一致性引擎 (全视频同产品同颜色同包装)',
    'UGC质量评分系统 (AI感/真人感/UGC感/广告转化 4维评分)',
    '运行性能Dashboard (实时任务状态)',
    '30天微信+远程技术支持',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Enterprise
doc.add_heading('企业版 — ¥4,999/月', level=2)
p = doc.add_paragraph()
p.add_run('适合: ').bold = True
p.add_run('批量运营团队、跨境电商企业、品牌方')
items = [
    '专业版全部功能',
    '多账号管理系统 (支持多个API Key轮换)',
    '无限批量生产模式 (不限任务数量)',
    'Watch Mode 持续监控模式 (放入新素材自动处理)',
    '商业授权系统 (机器码绑定+授权码验证+防篡改)',
    '多角色数字人库 (美国/英国/澳大利亚, 7种预设声音)',
    '多国家/多语言支持 (13个国家: US/UK/JP/KR/BR/ID/TH/VN/PH/MX/SA/AE/CN)',
    'UGC Director 表演指导引擎 (逐镜头情绪/手势/语速/停顿)',
    '批量任务Dashboard (全部任务状态一览)',
    '优先技术支持 (4小时内响应)',
    '定制功能开发 (按需定制, 如新平台接入/特殊工作流)',
    'API额度监控与预警 (自动提醒余额不足)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# === TECH ===
doc.add_heading('四、技术栈', level=1)
doc.add_paragraph(
    'AI模型: GPT-5.5 (脚本) | GPT Image / DALL-E 3 (图片) | Seedance 2.0 (视频) | ElevenLabs (语音)'
)
doc.add_paragraph(
    '平台接入: OpenAI API | Anthropic Claude | Google Gemini | DeepSeek | 火山引擎 ARK | Runway | Kling'
)
doc.add_paragraph(
    '视频处理: FFmpeg (合成/转码/字幕) | ffprobe (分析) | Whisper (语音转文字) | Tesseract (OCR)'
)

# === DELIVERY ===
doc.add_heading('五、交付与付款', level=1)
p = doc.add_paragraph()
p.add_run('交付方式: ').bold = True
p.add_run('GitHub私有仓库授权 + 部署文档 + 视频操作教程 + 远程协助')
p = doc.add_paragraph()
p.add_run('付款方式: ').bold = True
p.add_run('银行转账 (对公/对私) / 支付宝 / 微信支付 / USDT (TRC20)')
p = doc.add_paragraph()
p.add_run('发票: ').bold = True
p.add_run('可开具增值税普通发票或专用发票 (电子/纸质)')
p = doc.add_paragraph()
p.add_run('免费试用: ').bold = True
p.add_run('提供3天免费试用 (基础版全部功能), 满意后再付款')
p = doc.add_paragraph()
p.add_run('续费优惠: ').bold = True
p.add_run('年付享8折优惠 | 一次性购买3年送1年')

# === CONTACT ===
doc.add_heading('六、联系方式', level=1)
p = doc.add_paragraph()
p.add_run('GitHub: ').bold = True
p.add_run('https://github.com/massielvasquez193-dot/CC-SEEDSANCE-SKILL-ZIDONGHUA')
p = doc.add_paragraph()
p.add_run('邮箱: ').bold = True
p.add_run('(请联系销售代表获取)')
p = doc.add_paragraph()
p.add_run('微信: ').bold = True
p.add_run('(请联系销售代表获取)')

# === FOOTER ===
doc.add_paragraph()
doc.add_paragraph()
f1 = doc.add_paragraph()
f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = f1.add_run('TikTok AI Factory Pro — 全自动AI视频生产工厂 | 让每个人都能批量生产TikTok爆款视频')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

f2 = doc.add_paragraph()
f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = f2.add_run(f'本报价单有效期至 {datetime.now().year}年12月31日 | 价格不含税 | 本公司保留最终解释权')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

doc.save('报价单_中文版.docx')
print('Done: 报价单_中文版.docx')
