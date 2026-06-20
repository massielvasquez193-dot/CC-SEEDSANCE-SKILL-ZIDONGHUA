"""Generate Word versions of client documents"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_docx(md_path, docx_path, title_text):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.size = Pt(11)

    content = open(md_path, encoding='utf-8').read()
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(line[2:], level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # H2
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

        # H3
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)

        # Table
        elif line.startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1
            if len(table_lines) >= 2:
                rows = [[c.strip() for c in row.split('|')[1:-1]] for row in table_lines]
                rows = [r for r in rows if not all(c.replace('-','').replace(' ','') == '' for c in r)]
                if rows:
                    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Table Grid')
                    for ri, row in enumerate(rows):
                        for ci, cell_text in enumerate(row):
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = cell_text
                            if ri == 0:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.bold = True

        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)

        # Bullet
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle bold markers
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

        # Bold paragraph
        elif line.startswith('**') or '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

        # Regular text
        elif line and not line.startswith('---') and not line.startswith('>'):
            p = doc.add_paragraph(line)

        i += 1

    doc.save(docx_path)
    print(f'Done: {docx_path}')

# Generate all 3
create_docx('CLIENT_INSTALL_GUIDE.md', '客户端安装指南.docx', '安装指南')
create_docx('CLIENT_USER_GUIDE.md', '客户端使用指南.docx', '使用指南')
create_docx('FAQ.md', '常见问题FAQ.docx', 'FAQ')
